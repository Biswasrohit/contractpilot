"""Extract text from Word (.docx) documents."""

import io

from docx import Document


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract all text from a .docx file.

    Extracts text from paragraphs and tables (contracts often have terms in tables).

    Args:
        file_bytes: Raw .docx file bytes.

    Returns:
        Full extracted text with section breaks.
    """
    doc = Document(io.BytesIO(file_bytes))

    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts)
